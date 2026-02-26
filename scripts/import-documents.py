#!/usr/bin/env python3
"""
Document Importer
Converts Word/PDF documents to Markdown and imports into knowledge base.
"""

import os
import sys
import argparse
from pathlib import Path
from datetime import datetime
import re

try:
    from docx import Document as DocxDocument
    from PyPDF2 import PdfReader
    import markdown
except ImportError:
    print("❌ Error: Required libraries not installed")
    print("   Run: pip install -r scripts/requirements.txt")
    sys.exit(1)

# Base directory
BASE_DIR = Path(__file__).parent.parent
DOCS_DIR = BASE_DIR / "business-knowledge" / "documents"

# Department mapping
DEPARTMENT_MAP = {
    'marketing': BASE_DIR / "business-knowledge" / "departments" / "marketing" / "documents",
    'sales': BASE_DIR / "business-knowledge" / "departments" / "sales" / "documents",
    'customer-service': BASE_DIR / "business-knowledge" / "departments" / "customer-service" / "documents",
    'crm': BASE_DIR / "business-knowledge" / "departments" / "crm" / "documents",
    'website': BASE_DIR / "business-knowledge" / "departments" / "website" / "documents",
    'shared': DOCS_DIR
}


def clean_filename(filename):
    """Convert filename to clean markdown filename"""
    # Remove extension
    name = Path(filename).stem
    # Replace spaces and special chars with hyphens
    name = re.sub(r'[^\w\s-]', '', name)
    name = re.sub(r'[-\s]+', '-', name)
    return name.lower() + '.md'


def read_docx(file_path):
    """Read Word document and convert to markdown"""
    print(f"📄 Reading Word document: {file_path.name}")

    doc = DocxDocument(file_path)
    markdown_content = []

    # Add header
    markdown_content.append(f"# {file_path.stem}\n")
    markdown_content.append(f"*Imported from Word document on {datetime.now().strftime('%Y-%m-%d')}*\n")
    markdown_content.append(f"*Original file: {file_path.name}*\n\n---\n\n")

    # Convert paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            markdown_content.append("\n")
            continue

        # Check if heading
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', ''))
            markdown_content.append(f"{'#' * level} {text}\n\n")
        # Check if list item
        elif para.style.name.startswith('List'):
            markdown_content.append(f"- {text}\n")
        else:
            markdown_content.append(f"{text}\n\n")

    # Convert tables
    for table in doc.tables:
        markdown_content.append("\n")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            markdown_content.append("| " + " | ".join(cells) + " |\n")
            if i == 0:
                markdown_content.append("| " + " | ".join(['---'] * len(cells)) + " |\n")
        markdown_content.append("\n")

    return ''.join(markdown_content)


def read_pdf(file_path):
    """Read PDF and convert to markdown"""
    print(f"📄 Reading PDF document: {file_path.name}")

    reader = PdfReader(file_path)
    markdown_content = []

    # Add header
    markdown_content.append(f"# {file_path.stem}\n")
    markdown_content.append(f"*Imported from PDF on {datetime.now().strftime('%Y-%m-%d')}*\n")
    markdown_content.append(f"*Original file: {file_path.name}*\n\n---\n\n")

    # Extract text from all pages
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            markdown_content.append(f"## Page {page_num}\n\n")
            markdown_content.append(f"{text}\n\n")

    return ''.join(markdown_content)


def import_document(source_path, department=None, output_name=None):
    """Import a document and convert to markdown"""

    source_path = Path(source_path)

    if not source_path.exists():
        print(f"❌ Error: File not found: {source_path}")
        return False

    # Determine output path
    if department and department in DEPARTMENT_MAP:
        output_dir = DEPARTMENT_MAP[department]
    else:
        output_dir = DOCS_DIR

    output_dir.mkdir(parents=True, exist_ok=True)

    # Determine output filename
    if output_name:
        output_file = output_dir / output_name
    else:
        output_file = output_dir / clean_filename(source_path.name)

    # Convert based on file type
    file_ext = source_path.suffix.lower()

    try:
        if file_ext == '.docx':
            markdown = read_docx(source_path)
        elif file_ext == '.pdf':
            markdown = read_pdf(source_path)
        elif file_ext == '.txt':
            print(f"📄 Reading text file: {source_path.name}")
            with open(source_path, 'r') as f:
                content = f.read()
            markdown = f"# {source_path.stem}\n\n{content}"
        else:
            print(f"❌ Unsupported file type: {file_ext}")
            print("   Supported: .docx, .pdf, .txt")
            return False

        # Write markdown file
        with open(output_file, 'w') as f:
            f.write(markdown)

        print(f"✅ Converted: {output_file}")

        # Update department memory file if applicable
        if department and department in DEPARTMENT_MAP:
            update_memory_file(department, output_file.name, source_path.stem)

        return True

    except Exception as e:
        print(f"❌ Error converting file: {e}")
        return False


def update_memory_file(department, doc_filename, doc_title):
    """Add reference to document in department memory file"""

    memory_file = BASE_DIR / "business-knowledge" / "departments" / department / "memory.md"

    if not memory_file.exists():
        print(f"   ⚠️  Memory file not found: {memory_file}")
        return

    # Read current content
    with open(memory_file, 'r') as f:
        content = f.read()

    # Add reference to Knowledge Base section
    doc_ref = f"- **{doc_title}**: [documents/{doc_filename}](documents/{doc_filename})"

    if doc_ref in content:
        print(f"   ℹ️  Document already referenced in memory file")
        return

    # Find Knowledge Base section
    kb_section = "### Important Documents"

    if kb_section in content:
        # Add to existing section
        lines = content.split('\n')
        for i, line in enumerate(lines):
            if kb_section in line:
                # Insert after this line
                lines.insert(i + 2, doc_ref)
                break

        content = '\n'.join(lines)
    else:
        # Add new section
        content += f"\n\n### Important Documents\n{doc_ref}\n"

    # Write back
    with open(memory_file, 'w') as f:
        f.write(content)

    print(f"   ✅ Added reference to {department}/memory.md")


def import_directory(source_dir, department=None):
    """Import all supported documents from a directory"""

    source_dir = Path(source_dir)

    if not source_dir.exists():
        print(f"❌ Error: Directory not found: {source_dir}")
        return

    supported_extensions = ['.docx', '.pdf', '.txt']
    files = []

    for ext in supported_extensions:
        files.extend(source_dir.glob(f'*{ext}'))

    if not files:
        print(f"⚠️  No supported files found in {source_dir}")
        return

    print(f"📁 Found {len(files)} documents to import\n")

    success_count = 0
    for file_path in files:
        if import_document(file_path, department):
            success_count += 1
        print()  # Blank line between files

    print(f"✅ Successfully imported {success_count}/{len(files)} documents")


def main():
    """Main execution"""
    parser = argparse.ArgumentParser(
        description='Import Word/PDF documents into knowledge base',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Import single file to shared documents
  python scripts/import-documents.py Brand_Guidelines.docx

  # Import to specific department
  python scripts/import-documents.py Sales_Playbook.docx --department sales

  # Import all files from directory
  python scripts/import-documents.py ~/Documents/YNA_Docs/ --directory

  # Import to marketing with custom name
  python scripts/import-documents.py guide.docx --department marketing --output brand-guidelines.md
        """
    )

    parser.add_argument('source', help='Source file or directory path')
    parser.add_argument('--department', '-d',
                        choices=['marketing', 'sales', 'customer-service', 'crm', 'website', 'shared'],
                        help='Department to import to (default: shared)')
    parser.add_argument('--output', '-o', help='Output filename (default: auto-generated)')
    parser.add_argument('--directory', action='store_true',
                        help='Import all files from source directory')

    args = parser.parse_args()

    print("=" * 60)
    print("Document Importer")
    print("=" * 60 + "\n")

    if args.directory:
        import_directory(args.source, args.department)
    else:
        import_document(args.source, args.department, args.output)

    print("\n" + "=" * 60)


if __name__ == "__main__":
    main()
